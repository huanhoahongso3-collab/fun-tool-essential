#!/usr/bin/env python3

import argparse
import pyperclip
from docx import Document
import os

def read_input_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read().strip()
    elif ext == '.docx':
        doc = Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs])
    else:
        raise ValueError("Unsupported file format. Only .txt and .docx are supported.")

def write_output_file(file_path, content):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.txt':
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
    elif ext == '.docx':
        doc = Document()
        for line in content.splitlines():
            doc.add_paragraph(line)
        doc.save(file_path)
    else:
        raise ValueError("Unsupported output format. Only .txt and .docx are supported.")

def main():
    parser = argparse.ArgumentParser(description="thlinh: A tool for duplicating sentence/phrase/word")
    parser.add_argument('pos_input', nargs='?', help='Input string (positional alternative to -i if quoted)')
    parser.add_argument('-i', '--input', type=str, help="Input string to duplicate")
    parser.add_argument('-stt', action='store_true', help="Add rank numbering (1. 2. ...)")
    parser.add_argument('-t', '--times', type=int, default=1, help="Number of repetitions")
    parser.add_argument('-c', '--copy', action='store_true', help="Copy output to clipboard")
    parser.add_argument('-o', '--output', type=str, help="Output to .txt or .docx file")
    parser.add_argument('-nt', '--no-terminal', action='store_true', help="Suppress terminal output")
    parser.add_argument('-f', '--file', type=str, help="Input file (.txt or .docx)")

    args = parser.parse_args()

    # Input resolution
    if args.file:
        if args.input or args.pos_input or args.stt or args.copy:
            parser.error("When using -f, -i, positional input, -stt, and -c are not allowed.")
        args.no_terminal = True
        input_text = read_input_file(args.file)
    else:
        input_text = args.input or args.pos_input
        if not input_text:
            parser.error("No input provided. Use -i or quoted positional input (e.g. thlinh \"hello world\").")

    # Generate output
    output_lines = []
    for i in range(1, args.times + 1):
        line = f"{i}. {input_text}" if args.stt else input_text
        output_lines.append(line)
    output_content = '\n'.join(output_lines)

    if not args.no_terminal:
        print(output_content)

    if args.copy:
        pyperclip.copy(output_content)

    if args.output:
        write_output_file(args.output, output_content)

if __name__ == "__main__":
    main()
